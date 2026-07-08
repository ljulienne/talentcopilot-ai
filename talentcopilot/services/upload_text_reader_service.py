from dataclasses import dataclass
from pathlib import Path
from tempfile import NamedTemporaryFile


@dataclass
class UploadedTextDocument:
    filename: str
    file_type: str
    text: str
    status: str = "OK"


class UploadTextReaderService:
    SUPPORTED_EXTENSIONS = {".txt", ".pdf", ".docx"}

    def read_uploaded_file(self, uploaded_file) -> UploadedTextDocument:
        filename = getattr(uploaded_file, "name", "uploaded.txt")
        suffix = Path(filename).suffix.lower()
        file_type = suffix.replace(".", "") or "unknown"

        if suffix not in self.SUPPORTED_EXTENSIONS:
            return UploadedTextDocument(filename, file_type, "", f"Unsupported file type: {suffix}")

        try:
            raw = uploaded_file.getvalue()

            if suffix == ".txt":
                text = raw.decode("utf-8", errors="ignore") if isinstance(raw, bytes) else str(raw)
                return UploadedTextDocument(filename, "txt", text)

            with NamedTemporaryFile(delete=True, suffix=suffix) as tmp:
                tmp.write(raw)
                tmp.flush()

                if suffix == ".pdf":
                    return self._read_pdf(filename, tmp.name)

                if suffix == ".docx":
                    return self._read_docx(filename, tmp.name)

        except Exception as exc:
            return UploadedTextDocument(filename, file_type, "", f"Could not read file: {exc}")

        return UploadedTextDocument(filename, file_type, "", "Could not read file")

    def _read_pdf(self, filename: str, path: str) -> UploadedTextDocument:
        try:
            from pypdf import PdfReader
        except Exception:
            return UploadedTextDocument(
                filename=filename,
                file_type="pdf",
                text="",
                status="Missing dependency: pypdf. Add pypdf to requirements.txt and reboot Streamlit.",
            )

        try:
            reader = PdfReader(path)
            text = "\n".join((page.extract_text() or "") for page in reader.pages).strip()

            if not text:
                return UploadedTextDocument(
                    filename=filename,
                    file_type="pdf",
                    text="",
                    status="No extractable text found. This PDF may be scanned; OCR is not available yet.",
                )

            return UploadedTextDocument(filename=filename, file_type="pdf", text=text)

        except Exception as exc:
            return UploadedTextDocument(filename, "pdf", "", f"Could not parse PDF: {exc}")

    def _read_docx(self, filename: str, path: str) -> UploadedTextDocument:
        try:
            import docx
        except Exception:
            return UploadedTextDocument(
                filename=filename,
                file_type="docx",
                text="",
                status="Missing dependency: python-docx. Add python-docx to requirements.txt and reboot Streamlit.",
            )

        try:
            document = docx.Document(path)
            text = "\n".join(paragraph.text for paragraph in document.paragraphs).strip()

            if not text:
                return UploadedTextDocument(filename, "docx", "", "No extractable text found in DOCX.")

            return UploadedTextDocument(filename, "docx", text)

        except Exception as exc:
            return UploadedTextDocument(filename, "docx", "", f"Could not parse DOCX: {exc}")

    def read_text_direct(self, filename: str, text: str) -> UploadedTextDocument:
        suffix = Path(filename).suffix.lower().replace(".", "") or "txt"
        return UploadedTextDocument(filename=filename, file_type=suffix, text=text)
