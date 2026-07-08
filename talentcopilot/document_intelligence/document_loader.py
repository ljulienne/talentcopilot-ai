from pathlib import Path

from talentcopilot.document_intelligence.models import LoadedDocument


class DocumentLoader:
    def load_text(self, filename: str, content: str) -> LoadedDocument:
        suffix = Path(filename).suffix.lower().replace(".", "") or "txt"
        return LoadedDocument(filename=filename, file_type=suffix, text=content)

    def load_path(self, path: str) -> LoadedDocument:
        p = Path(path)
        suffix = p.suffix.lower()

        if suffix == ".txt":
            return LoadedDocument(filename=p.name, file_type="txt", text=p.read_text(encoding="utf-8"))

        if suffix == ".pdf":
            return LoadedDocument(filename=p.name, file_type="pdf", text=self._load_pdf(p))

        if suffix == ".docx":
            return LoadedDocument(filename=p.name, file_type="docx", text=self._load_docx(p))

        raise ValueError(f"Unsupported document type: {suffix}")

    def _load_pdf(self, path: Path) -> str:
        try:
            from pypdf import PdfReader
        except Exception as exc:
            raise ImportError("pypdf is required to load PDF files. Install with: pip install pypdf") from exc

        reader = PdfReader(str(path))
        parts = []
        for page in reader.pages:
            parts.append(page.extract_text() or "")
        return "\n".join(parts).strip()

    def _load_docx(self, path: Path) -> str:
        try:
            import docx
        except Exception as exc:
            raise ImportError("python-docx is required to load DOCX files. Install with: pip install python-docx") from exc

        document = docx.Document(str(path))
        return "\n".join(paragraph.text for paragraph in document.paragraphs).strip()
